
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile

def set_paragraph_style(paragraph, font_size=11, bold=False):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.size = Pt(font_size)
    font.bold = bold
    paragraph.paragraph_format.line_spacing = 1.4

def create_application_docx(current_key, result, requirements, selections, output2_text_list, file_path):
    doc = Document()
    doc.add_heading("의약품 허가 후 제조방법 변경관리 가이드라인(민원인 안내서)", level=0)
    doc.styles['Normal'].font.name = 'NanumGothic'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    # 1. 신청인
    doc.add_heading("1. 신청인", level=1)
    table1 = doc.add_table(rows=3, cols=2)
    table1.style = "Table Grid"
    for i, label in enumerate(["성명", "제조소(영업소) 명칭", "변경신청 제품명"]):
        table1.cell(i, 0).text = label
        table1.cell(i, 1).text = ""

    # 2. 변경유형
    doc.add_heading("2. 변경유형", level=1)
    t2 = doc.add_table(rows=1, cols=1)
    t2.style = "Table Grid"
    t2.cell(0, 0).text = result.get("title_text", "")

    # 3. 신청유형
    doc.add_heading("3. 신청유형", level=1)
    t3 = doc.add_table(rows=2, cols=2)
    t3.style = "Table Grid"
    t3.cell(0, 0).text = "분류"
    t3.cell(0, 1).text = result.get("output_1_tag", "")
    t3.cell(1, 0).merge(t3.cell(1, 1)).text = result.get("output_1_text", "")

    # 4. 충족조건
    doc.add_heading("4. 충족조건", level=1)
    t4 = doc.add_table(rows=1, cols=2)
    t4.style = "Table Grid"
    t4.cell(0, 0).text = "충족조건"
    t4.cell(0, 1).text = "조건 충족 여부"
    for rk, text in requirements.items():
        row = t4.add_row().cells
        row[0].text = text
        sel = selections.get(f"{current_key}_req_{rk}", "")
        row[1].text = "○" if sel == "충족" else "×" if sel == "미충족" else ""

    # 5. 필요서류
    doc.add_heading("5. 필요서류", level=1)
    t5 = doc.add_table(rows=1, cols=1)
    t5.style = "Table Grid"
    t5.cell(0, 0).text = "서류"
    for line in output2_text_list:
        row = t5.add_row().cells
        row[0].text = line.strip()

    doc.save(file_path)

if "step" not in st.session_state:
    st.session_state.step = 1
if st.session_state.step == 8:
    step7_results = st.session_state.get("step7_results", {})
    step6_items = st.session_state.get("step6_items", {})
    step6_selections = st.session_state.get("step6_selections", {})

    title_keys = list(step7_results.keys())
    if "step8_page" not in st.session_state:
        st.session_state.step8_page = 0

    page = st.session_state.step8_page
    total_pages = len(title_keys)
    current_key = title_keys[page]
    result = step7_results[current_key]
    requirements = step6_items.get(current_key, {}).get("requirements", {})
    selections = {
        f"{current_key}_req_{rk}": step6_selections.get(f"{current_key}_req_{rk}", "")
        for rk in requirements
    }
    output2_text_list = [line.strip() for line in result.get("output_2_text", "").split("\n") if line.strip()]

    col1, col2, col3 = st.columns([1, 3, 1])
    with col1:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            create_application_docx(current_key, result, requirements, selections, output2_text_list, tmp.name)
            with open(tmp.name, "rb") as f:
                st.download_button("📄 파일 다운로드", f, file_name=f"신청서_{current_key}.docx")

    with col2:
        st.markdown(
            f"<h5 style='text-align:center'>「의약품 허가 후 제조방법 변경관리 가이드라인(민원인 안내서)」[붙임] 신청양식 예시<br>{page+1} / {total_pages}</h5>",
            unsafe_allow_html=True,
        )

    with col3:
        if st.button("🖨 인쇄하기"):
            st.markdown("<script>window.print();</script>", unsafe_allow_html=True)

    html = f"""
    <style>
    table, th, td {{ border: 1px solid black; border-collapse: collapse; padding: 6px; text-align: center; }}
    th, td {{ font-size: 14px; }}
    </style>

    <h5>1. 신청인</h5>
    <table><tr><td>성명</td><td></td></tr><tr><td>제조소(영업소) 명칭</td><td></td></tr><tr><td>변경신청 제품명</td><td></td></tr></table><br>

    <h5>2. 변경유형</h5>
    <table><tr><td>{result["title_text"]}</td></tr></table><br>

    <h5>3. 신청유형</h5>
    <table><tr><td>분류</td><td>{result["output_1_tag"]}</td></tr>
    <tr><td colspan="2">{result["output_1_text"].replace("\n", "<br>")}</td></tr></table><br>

    <h5>4. 충족조건</h5>
    <table><tr><th>충족조건</th><th>조건 충족 여부</th></tr>
    """
    for rk, text in requirements.items():
        state = selections[f"{current_key}_req_{rk}"]
        symbol = "○" if state == "충족" else "×" if state == "미충족" else ""
        html += f"<tr><td style='text-align:left'>{text}</td><td>{symbol}</td></tr>"
    html += "</table><br>"

    html += "<h5>5. 필요서류</h5><table><tr><th>서류</th></tr>"
    for line in output2_text_list:
        html += f"<tr><td style='text-align:left'>{line.strip()}</td></tr>"
    html += "</table><br>"

    st.markdown(html, unsafe_allow_html=True)

    col_left, col_right = st.columns(2)
    with col_left:
        if st.button("⬅ 이전"):
            if st.session_state.step8_page == 0:
                st.session_state.step = 7
                del st.session_state["step8_page"]
            else:
                st.session_state.step8_page -= 1

    with col_right:
        if st.button("다음 ➡") and st.session_state.step8_page < total_pages - 1:
            st.session_state.step8_page += 1
